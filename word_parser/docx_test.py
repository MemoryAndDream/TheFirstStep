# -*- coding: utf-8 -*-
# @Time    : 2018/10/29 22:42
# @Author  : meng_zhihao
# @Email   : 312141830@qq.com
# @File    : docx_test.py
import docx
import sys
import pickle
import re
import  codecs
import string
import shutil
from win32com import client as wc

def docx_try():

    doc=docx.Document(r'C:\Users\ASUS\Desktop\3.docx')
    for p in doc.paragraphs:

        print(p.text)

    for t in doc.tables:

        for r in t.rows:

            for c in r.cells:

                print(c.text)

def convert():
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(r'C:\Users\ASUS\Desktop\3.doc')        # 目标路径下的文件
    doc.SaveAs(r'C:\Users\ASUS\Desktop\3.docx', 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()


if __name__ == '__main__':
#    convert()
    docx_try()