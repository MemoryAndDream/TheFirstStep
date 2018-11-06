# -*- coding: utf-8 -*-
# @Time    : 2018/10/29 22:42
# @Author  : meng_zhihao
# @Email   : 312141830@qq.com
# @File    : docx_test.py
import docx
import sys
import re
import os
from win32com import client as wc
reload(sys)
sys.setdefaultencoding('utf8')
import  xlwt
import datetime
import ConfigParser
import jieba
content = open('main.conf').read()
#Window下用记事本打开配置文件并修改保存后，编码为UNICODE或UTF-8的文件的文件头
#会被相应的加上\xff\xfe（\xff\xfe）或\xef\xbb\xbf，然后再传递给ConfigParser解析的时候会出错
#，因此解析之前，先替换掉
content = re.sub(r"\xfe\xff","", content)
content = re.sub(r"\xff\xfe","", content)
content = re.sub(r"\xef\xbb\xbf","", content)
open('main.conf', 'w').write(content)
mainconf = ConfigParser.ConfigParser()
mainconf.read('main.conf')

xls_FILE_NAME =unicode(mainconf.get('conf','xls_FILE_NAME'))
DATA_DIR = unicode(mainconf.get('conf','DATA_DIR'))
TMP_DOCX =  unicode(mainconf.get('conf','TMP_DOCX'))
CACHE_TXT = unicode(mainconf.get('conf','CACHE_TXT'))

print xls_FILE_NAME,DATA_DIR,TMP_DOCX,CACHE_TXT


#字段名称
CSV_TITLE = [
    "证券代码",
    "证券简称",
    "投资者关系活动类别",
    "参与单位名称",
    "参与人员姓名",
    "参与单位名称及人员姓名",
    "时间",
    "地点",
    "上市公司接待人员级别",
    "上市公司接待人员姓名", #姓名没有规律，提取不了
    "上市公司接待人员",
    "投资者关系活动主要内容介绍",
    "附件清单",
    "日期",
    "披露日期",
    "参与单位机构个数",
    "参与人员人数",
    "接待人员人数",
    "投资者问题个数",
    "投资者问题字数",
    "公司回答字数",
   # "原文件路径"
]

#机构核心词 通过分词统计得到  包含这些词的视为机构
ORG_WORDS = [
    u'基金',u'证券',u'券商',u'个人',u'分析',u'评级',u'投资',u'公司',u'银行',u'单位',u'摩根',u'媒体'
]
RANK_WORDS = [u'董事长',u'副总经理',u'总经理',u'秘书',u'董事',u'财务总监',u'副总裁',u'总裁',u'总监',u'经理',u'高管',u'助理',u'级别']

BLACK_LIST = [u'研究员',u'等'] # 不属于人名也不属于职位

family_names  = [
'赵', '钱', '孙', '李', '周', '吴', '郑', '王', '冯', '陈', '褚', '卫', '蒋', '沈', '韩', '杨', '朱', '秦', '尤', '许',
'何', '吕', '施', '张', '孔', '曹', '严', '华', '金', '魏', '陶', '姜', '戚', '谢', '邹', '喻', '柏', '水', '窦', '章',
'云', '苏', '潘', '葛', '奚', '范', '彭', '郎', '鲁', '韦', '昌', '马', '苗', '凤', '花', '方', '俞', '任', '袁', '柳',
'酆', '鲍', '史', '唐', '费', '廉', '岑', '薛', '雷', '贺', '倪', '汤', '滕', '殷', '罗', '毕', '郝', '邬', '安', '常',
'乐', '于', '时', '傅', '皮', '卞', '齐', '康', '伍', '余', '元', '卜', '顾', '孟', '平', '黄', '和', '穆', '萧', '尹',
'姚', '邵', '堪', '汪', '祁', '毛', '禹', '狄', '米', '贝', '明', '臧', '计', '伏', '成', '戴', '谈', '宋', '茅', '庞',
'熊', '纪', '舒', '屈', '项', '祝', '董', '梁', '邓']

family_names = [family_name.decode('utf8') for family_name in family_names]

book = xlwt.Workbook()
#创建表单
sheet1 = book.add_sheet(u'sheet1',cell_overwrite_ok=True)

# docx文件解析
def parser_docx(doc_path=r'C:\Users\Administrator\Desktop\mytest\61937348.DOCX'):
    '''
    docx文件解析
    :param doc_path:doc文件路径
    :return:
    '''
    record_obj = {}
    for title in CSV_TITLE:
        record_obj[title]=''
    doc=docx.Document(doc_path)
    for p in doc.paragraphs:
        text = p.text
        # 证券代码解析
        zqdm = get_regex_data(u'证券代码：\s*(\d+)',text)
        if zqdm:
           print zqdm
           record_obj['证券代码'] = zqdm
        # 证券简称解析
        zqjc = get_regex_data(u"证券简称：([\u4E00-\u9FA5]*)",text)
        if zqjc:
            print zqjc
            record_obj['证券简称'] = zqjc
    for t in doc.tables:
        for row in t.rows:
            if len(row.cells)<2:continue
            cells = row.cells
            cells[0].text = cells[0].text.replace('\n','').replace('\r','')
            # 关系活动类别解析
            if u'关系活动类别' in cells[0].text:
                for activity_type in [u'□特定对象调研',u'□分析师会议',u'□媒体采访',u'□业绩说明会',u'□新闻发布会',u'□路演活动',u'□现场参观',u'□其他']:
                    if activity_type not in cells[1].text:
                        activity_type = activity_type.replace(u'□',u'')
                        record_obj['投资者关系活动类别'] = activity_type
                        print activity_type
                        break
            # 参与单位名称及人员姓名解析 不规则字段，解析规则：大于3个字的或者包含关键词的作为机构名 其他作为人名
            elif u'参与单位名称及人员姓名' in cells[0].text:
                content = cells[1].text
                record_obj["参与单位名称及人员姓名"]=content
                geren = []
                jigou = []
                names = re.findall(u'[\u4E00-\u9FA5]+',content)
                for name in names:
                    is_ORG = False
                    if name in BLACK_LIST:
                        continue
                    for word in ORG_WORDS:
                        if word in name:
                            is_ORG = True
                            break
                    if len(name) > 3:  # 四个字的也当成机构名
                        is_ORG = True
                    if is_ORG:
                        flag = True
                        for family_name in family_names:
                            if family_name in name:
                                seg_list = [w for w in jieba.cut(name)]
                                if seg_list[-1][0] in family_names:
                                    geren.append(seg_list[-1])
                                    jigou.append(u''.join(seg_list[:-1]))
                                    flag = False
                                elif len(seg_list)>1 and seg_list[-2] in family_names:
                                    geren.append(seg_list[-2]+seg_list[-1])
                                    jigou.append(u''.join(seg_list[:-1]))
                                    flag = False
                                break
                        if flag:jigou.append(name)
                    else:
                        geren.append(name)
                jigou = list(set(jigou))
                record_obj['参与单位名称'] = jigou
                record_obj['参与人员姓名'] = geren
                record_obj['参与单位机构个数'] = len(jigou)
                record_obj['参与人员人数'] = len(geren)
                print names
            elif u'时间' in cells[0].text:
                record_obj['时间'] = cells[1].text
            elif u'日期' in cells[0].text:
                record_obj['日期'] = cells[1].text
            elif u'地点' in cells[0].text:
                record_obj['地点'] = cells[1].text
            elif u'上市公司接待人员' in cells[0].text:
                record_obj['上市公司接待人员'] = cells[1].text
            elif u'投资者关系活动主要内容' in cells[0].text:
                record_obj["投资者关系活动主要内容介绍"] = cells[1].text
            elif u'附件清单'  in cells[0].text:
                record_obj["附件清单"] = cells[1].text
        if record_obj['上市公司接待人员']:
            jdry = record_obj['上市公司接待人员'] # 不规则字段 解析规则
            renmin = []
            jibie = []
            names = re.findall(u'[\u4E00-\u9FA5]+', jdry)
            for name in names:
                is_RANK = False
                for word in RANK_WORDS:
                    if word in name:
                        is_RANK = True
                        break
                if len(name) > 3:  # 四个字的也当成机构名
                    is_RANK = True

                if is_RANK:
                    for family_name in family_names:
                        if family_name in name:
                            seg_list = [w for w in jieba.cut(name)]
                            if seg_list[-1][0] in family_names:
                                renmin.append(seg_list[-1])
                                jibie.append(u''.join(seg_list[:-1]))
                                flag = False
                            elif len(seg_list) > 1 and seg_list[-2] in family_names:
                                renmin.append(seg_list[-2]+seg_list[-1])
                                jibie.append(u''.join(seg_list[:-1]))
                                flag = False
                            break
                    if flag:jibie.append(name)
                else:
                    renmin.append(name)
            jibie = list(set(jibie))
            record_obj['上市公司接待人员级别'] =  u' '.join(jibie)
            record_obj['上市公司接待人员姓名'] = renmin
        if record_obj['参与人员姓名']:
            record_obj['参与人员人数'] = len(record_obj['参与人员姓名'])
            record_obj['参与人员姓名'] = u' '.join(record_obj['参与人员姓名'])
        if record_obj['上市公司接待人员姓名']:
            record_obj['接待人员人数'] = len(record_obj['上市公司接待人员姓名'])
            record_obj['上市公司接待人员姓名'] = u' '.join(record_obj['上市公司接待人员姓名'])
        if record_obj["投资者关系活动主要内容介绍"]:
            nrjs = record_obj["投资者关系活动主要内容介绍"]
            lines = nrjs.split('\n')
            questions = 0
            answers = 0
            questions_words_num = 0
            answers_words_num = 0
            '''放弃了的逻辑 因为需求建议只解析标准格式的 就是说问题尽量往少里统计'''
            # for line in lines:
            #     line = line.strip()
            #     if not line:continue
            #     if re.match(u'.*[\?？]$',line):
            #         questions += 1
            #         questions_words_num += len(line)
            #     elif len(line) >28:
            #         answers_words_num += len(line)
            #     else:
            #         if re.match(u'^\d[,，、\s]',line):
            #             questions += 1
            #             questions_words_num += len(line)
            #             #print line,'q'
            #         elif not re.match(u'.*[:.。：%]$',line): #可以加上答的判断
            #             questions += 1
            #             questions_words_num += len(line)
            #         else:
            #             answers_words_num+= len(line)
            '''新的逻辑只检测 xx ： 或者 xxx？ 作为问题和回答的判断依据 所以问题的格式和字数都会偏小'''
            label_list = ['a' for i in range(len(lines)) ]
            a =q =0
            for i in range(len(lines)):
                line = lines[i]
                line = line.strip()
                if not line:continue
                if re.match(u'.*[\?？]$', line):
                    label_list[i]='q'
                    a+=1
                elif re.match(u'^\s*[\u4E00-\u9FA5]{0,4}答[\u4E00-\u9FA5]{0,4}\s*：', line):
                    label_list[i] = 'a'
                    if i >1:label_list[i-1] = 'q'
                    q+=1
            for i in range(len(lines)):
                if label_list[i] == 'q':
                    questions += 1
                    questions_words_num += len(lines[i])
                else:
                    answers += 1
                    answers_words_num += len(lines[i])
            record_obj["公司回答字数"] = answers_words_num
            record_obj["投资者问题个数"] = max(a,q)
            record_obj["投资者问题字数"] = questions_words_num
    return record_obj


def doc_to_docx(doc_path,docx_path):
    '''
    把doc文件转成docx文件方便解析
    :return:
    '''
   # import pdb;pdb.set_trace()
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(doc_path)        # 目标路径下的文件
    doc.SaveAs(docx_path, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()

def get_regex_data(regex, content):
    '''
    正则处理方法
    '''
    group = re.search(regex, content)
    if group:
        return group.groups()[0]
    else:
        return ''



def main(): #
    '''
    遍历文件目录下所有有效文件 doc docx pdf文件进行一一处理

    :return:
    '''
    file_path = ''
   # create_csv()
    file_list = os.walk(DATA_DIR)
    readed_file_name_list = []
    if not os.path.exists(CACHE_TXT):
        with open(CACHE_TXT,'w') as f:
            f.write(u'已读取的文件列表')
    with open(CACHE_TXT) as f:
        for line in f:
            readed_file_name_list.append(line.strip())

    tmp_xls_FILE_NAME = datetime.datetime.now().strftime('%Y%m%d%H%M%S') + xls_FILE_NAME
    i = j = 0
    for field in CSV_TITLE:
        sheet1.write(i, j, unicode(field))
        j += 1
    i = 1
    for file_paths in file_list:
        sub_file_paths = file_paths[2]
        parent_path = file_paths[0]
        for file_path in sub_file_paths:
            try:
                if file_path in readed_file_name_list:
                    continue
                if '$' in file_path or file_path in TMP_DOCX: #缓存文件
                    continue
                lower_file_path = file_path.lower()
                if  lower_file_path.endswith(u'doc')  or lower_file_path.endswith(u'docx') : #or lower_file_path.endswith(u'pdf')
                    tmpPath = os.path.join(parent_path, file_path)
                    if lower_file_path.endswith(u'doc'):
                        doc_to_docx(tmpPath,TMP_DOCX)
                        tmpPath = TMP_DOCX
                    obj = parser_docx(tmpPath)
                    csv_row = []
                    j=0
                    for field in CSV_TITLE:
                        sheet1.write(i, j, obj[field]) # 这里有问题
                        j += 1
                    i+=1
                    # 保存文件
                    book.save(tmp_xls_FILE_NAME)
                else:
                    pass
                with open(CACHE_TXT,'a') as f:
                    f.write(file_path+'\n')
            except Exception,e:
                print str(e)




# 打包命令 pyinstaller -F -p D:\aliyun_meng\taobao_requests\ D:\aliyun_meng\taobao_requests\doc_parser.py
if __name__ == '__main__':
    main()
    # rs= parser_docx(doc_path=r'C:\Users\Administrator\Desktop\mytest\1.docx')
    # for o in rs:
    #    print o,rs[o]

